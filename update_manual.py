#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新NetOps使用手册内容
"""

import os
import sys
from pathlib import Path
from typing import Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装python-docx库")
    print("请运行: pip install python-docx")
    sys.exit(1)


class ManualUpdater:
    """手册更新器类"""

    def __init__(self, doc_path: str):
        """初始化

        Args:
            doc_path: 文档路径
        """
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"文档不存在: {doc_path}")

        self.doc = Document(str(self.doc_path))

    def add_paragraph(self, text: str, bold: bool = False, size: int = 11,
                     color: Optional[Tuple[int, int, int]] = None,
                     alignment: Optional[int] = None, space_after: int = 6):
        """添加段落

        Args:
            text: 段落文本
            bold: 是否加粗
            size: 字体大小
            color: 字体颜色（RGB元组）
            alignment: 对齐方式
            space_after: 段后间距

        Returns:
            段落对象
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_bullet(self, text: str, level: int = 0):
        """添加项目符号

        Args:
            text: 文本内容
            level: 缩进级别

        Returns:
            段落对象
        """
        p = self.doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
        if level > 0:
            p.paragraph_format.left_indent = Cm(1.5 * level)
        return p

    def add_heading(self, text: str, level: int = 1):
        """添加标题

        Args:
            text: 标题文本
            level: 标题级别（1-9）

        Returns:
            标题对象
        """
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        return heading

    def add_tip_box(self, text: str):
        """添加提示框

        Args:
            text: 提示内容
        """
        p = self.doc.add_paragraph()
        run = p.add_run('💡 提示：')
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(22, 93, 255)

        run2 = p.add_run(text)
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run2.font.size = Pt(10)

        p.paragraph_format.space_after = Pt(6)
        return p

    def add_warning_box(self, text: str):
        """添加警告框

        Args:
            text: 警告内容
        """
        p = self.doc.add_paragraph()
        run = p.add_run('⚠️ 警告：')
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 140, 0)

        run2 = p.add_run(text)
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run2.font.size = Pt(10)

        p.paragraph_format.space_after = Pt(6)
        return p

    def save(self, output_path: Optional[str] = None):
        """保存文档

        Args:
            output_path: 输出路径（可选，默认覆盖原文件）
        """
        save_path = Path(output_path) if output_path else self.doc_path

        try:
            self.doc.save(str(save_path))
            print(f'✅ 文档已保存: {save_path}')
        except Exception as e:
            print(f'❌ 保存失败: {str(e)}')
            sys.exit(1)

    def get_document_info(self) -> dict:
        """获取文档信息

        Returns:
            文档信息字典
        """
        return {
            'path': str(self.doc_path),
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections),
        }


def main():
    """主函数"""
    # 文档路径
    doc_path = r'E:\LobsterAI\Network-Config\NetOps_V1.0_使用说明书.docx'

    try:
        # 创建更新器
        updater = ManualUpdater(doc_path)

        # 打印文档信息
        info = updater.get_document_info()
        print(f'文档信息:')
        print(f'  - 路径: {info["path"]}')
        print(f'  - 段落数: {info["paragraphs"]}')
        print(f'  - 表格数: {info["tables"]}')
        print(f'  - 节数: {info["sections"]}')

        # 示例：添加新功能说明
        updater.add_heading('新功能特性', level=2)
        updater.add_bullet('支持多厂商设备配置生成')
        updater.add_bullet('新增AI智能分析功能')
        updater.add_bullet('优化运维工具箱性能')

        # 添加提示框
        updater.add_tip_box('建议定期更新设备清单，确保配置准确性。')

        # 保存文档
        updater.save()

    except Exception as e:
        print(f'❌ 错误: {str(e)}')
        sys.exit(1)


if __name__ == '__main__':
    main()
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(78, 89, 105)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_warning_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('⚠ 注意：')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(245, 63, 63)
    run2 = p.add_run(text)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(78, 89, 105)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows, col_widths=None):
    from docx.table import Table
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        shading_elm.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '165DFF')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F7FA')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

# Find the "常见问题" section and insert the security section before it
# We need to find the paragraph that contains "5. 常见问题"
target_index = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '5. 常见问题':
        target_index = i
        break

if target_index is None:
    # Fallback: find any paragraph containing 常见问题
    for i, para in enumerate(doc.paragraphs):
        if '常见问题' in para.text:
            target_index = i
            break

if target_index:
    # Get the element before which we need to insert
    target_element = doc.paragraphs[target_index]._element
    parent = target_element.getparent()

    # We need to insert a page break and the new section before the 常见问题 section
    # Since python-docx doesn't support inserting at arbitrary positions easily,
    # we'll insert the new elements before the target element in the XML

    def make_page_break():
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        return p

    def make_paragraph(text, bold=False, size=11, color=None, alignment=None, space_after=6):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), str(int(space_after * 20)))
        pPr.append(spacing)
        if alignment is not None:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), {WD_ALIGN_PARAGRAPH.CENTER: 'center', WD_ALIGN_PARAGRAPH.LEFT: 'left'}.get(alignment, 'left'))
            pPr.append(jc)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
            rPr.append(c)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_heading(text, level=2):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'Heading{level}')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_bullet(text, level=0):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'ListBullet')
        pPr.append(pStyle)
        if level > 0:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(int(720 * level)))
            pPr.append(ind)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_tip(text):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '80')
        pPr.append(spacing)
        p.append(pPr)

        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        rFonts1 = OxmlElement('w:rFonts')
        rFonts1.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts1.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr1.append(rFonts1)
        sz1 = OxmlElement('w:sz')
        sz1.set(qn('w:val'), '20')
        rPr1.append(sz1)
        b1 = OxmlElement('w:b')
        rPr1.append(b1)
        c1 = OxmlElement('w:color')
        c1.set(qn('w:val'), '165DFF')
        rPr1.append(c1)
        r1.append(rPr1)
        t1 = OxmlElement('w:t')
        t1.set(qn('xml:space'), 'preserve')
        t1.text = '💡 提示：'
        r1.append(t1)
        p.append(r1)

        r2 = OxmlElement('w:r')
        rPr2 = OxmlElement('w:rPr')
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts2.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr2.append(rFonts2)
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '20')
        rPr2.append(sz2)
        c2 = OxmlElement('w:color')
        c2.set(qn('w:val'), '4E5969')
        rPr2.append(c2)
        r2.append(rPr2)
        t2 = OxmlElement('w:t')
        t2.set(qn('xml:space'), 'preserve')
        t2.text = text
        r2.append(t2)
        p.append(r2)
        return p

    def make_warning(text):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '80')
        pPr.append(spacing)
        p.append(pPr)

        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        rFonts1 = OxmlElement('w:rFonts')
        rFonts1.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts1.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr1.append(rFonts1)
        sz1 = OxmlElement('w:sz')
        sz1.set(qn('w:val'), '20')
        rPr1.append(sz1)
        b1 = OxmlElement('w:b')
        rPr1.append(b1)
        c1 = OxmlElement('w:color')
        c1.set(qn('w:val'), 'F53F3F')
        rPr1.append(c1)
        r1.append(rPr1)
        t1 = OxmlElement('w:t')
        t1.set(qn('xml:space'), 'preserve')
        t1.text = '⚠ 注意：'
        r1.append(t1)
        p.append(r1)

        r2 = OxmlElement('w:r')
        rPr2 = OxmlElement('w:rPr')
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rFonts2.set(qn('w:ascii'), 'Microsoft YaHei')
        rPr2.append(rFonts2)
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '20')
        rPr2.append(sz2)
        c2 = OxmlElement('w:color')
        c2.set(qn('w:val'), '4E5969')
        rPr2.append(c2)
        r2.append(rPr2)
        t2 = OxmlElement('w:t')
        t2.set(qn('xml:space'), 'preserve')
        t2.text = text
        r2.append(t2)
        p.append(r2)
        return p

    def make_table_xml(headers, rows):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)
        tbl.append(tblPr)

        tblGrid = OxmlElement('w:tblGrid')
        for w in [2400, 6000, 7200]:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w))
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        # Header row
        tr = OxmlElement('w:tr')
        for h in headers:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '165DFF')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc_p = OxmlElement('w:jc')
            jc_p.set(qn('w:val'), 'center')
            pPr.append(jc_p)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
            b = OxmlElement('w:b')
            rPr.append(b)
            c = OxmlElement('w:color')
            c.set(qn('w:val'), 'FFFFFF')
            rPr.append(c)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = h
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

        # Data rows
        for ri, row in enumerate(rows):
            tr = OxmlElement('w:tr')
            for ci, val in enumerate(row):
                tc = OxmlElement('w:tc')
                if ri % 2 == 1:
                    tcPr = OxmlElement('w:tcPr')
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F5F7FA')
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)
                    tc.append(tcPr)
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')
                rPr.append(sz)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = str(val)
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        return tbl

    # Insert elements before the 常见问题 heading
    elements_to_insert = [
        make_page_break(),
        make_heading('4.4 设备密码加密存储（信息安全）', level=2),
        make_paragraph('本平台高度重视信息安全，所有设备登录密码均采用业界标准的加密方案进行存储，确保敏感数据不会以明文形式泄露。', size=11, space_after=8),

        make_heading('4.4.1 加密算法概述', level=3),
        make_paragraph('设备密码加密采用 AES-256-GCM 认证加密算法，这是目前业界公认最安全的对称加密方案之一，被广泛应用于金融、政府和军事领域。', size=11, space_after=6),

        make_table_xml(
            ['加密组件', '技术规格', '安全说明'],
            [
                ['加密算法', 'AES-256-GCM', '256位密钥，GCM模式提供认证加密，同时保证机密性和完整性'],
                ['密钥派生', 'PBKDF2-HMAC-SHA256', '基于密码的密钥派生函数，使用SHA-256哈希，迭代60万次，有效抵御暴力破解'],
                ['密钥来源', '本机唯一标识', '基于Windows MachineGuid + 主机名生成，每台计算机的密钥唯一'],
                ['随机数(Nonce)', '12字节随机生成', '每次加密使用不同的随机Nonce，相同密码每次加密结果不同'],
                ['存储格式', 'ENC: + Base64编码', '密文以 ENC: 前缀标识，Base64编码存储，便于识别和传输'],
                ['加密库', 'Python Cryptography', '使用Python官方推荐的cryptography库，经过广泛安全审计'],
            ]
        ),

        make_heading('4.4.2 加密流程', level=3),
        make_paragraph('密码加密与解密流程如下：', size=11, space_after=6),

        make_paragraph('加密过程（保存密码时）：', bold=True, size=11, space_after=4),
        make_bullet('① 采集本机唯一标识（Windows MachineGuid + 主机名）'),
        make_bullet('② 通过 PBKDF2-HMAC-SHA256 派生256位加密密钥（迭代60万次）'),
        make_bullet('③ 生成12字节随机 Nonce'),
        make_bullet('④ 使用 AES-256-GCM 加密明文密码'),
        make_bullet('⑤ 将 Nonce + 密文 进行 Base64 编码'),
        make_bullet('⑥ 添加 ENC: 前缀标识，写入设备清单文件'),

        make_paragraph('解密过程（使用密码时）：', bold=True, size=11, space_after=4),
        make_bullet('① 检测密码字段是否以 ENC: 开头'),
        make_bullet('② 从本机标识重新派生相同的加密密钥'),
        make_bullet('③ Base64解码，分离 Nonce 和密文'),
        make_bullet('④ 使用 AES-256-GCM 解密并验证完整性'),
        make_bullet('⑤ 返回明文密码供SSH/Telnet连接使用'),

        make_heading('4.4.3 安全特性', level=3),

        make_paragraph('🔐 本机绑定', bold=True, size=11, space_after=4),
        make_paragraph('加密密钥由本机唯一标识派生，即使他人获取了加密后的设备清单文件，也无法在其他计算机上解密。这确保了：', size=11, space_after=4),
        make_bullet('设备清单文件被拷贝到其他电脑后，密码无法解密'),
        make_bullet('每台计算机的加密密钥独立，互不通用'),
        make_bullet('有效防止文件泄露导致的密码外泄'),

        make_paragraph('🔐 认证加密', bold=True, size=11, space_after=4),
        make_paragraph('AES-GCM 模式不仅加密数据，还提供完整性校验。任何对密文的篡改都会导致解密失败并被检测到，防止密文被恶意修改。', size=11, space_after=6),

        make_paragraph('🔐 抗暴力破解', bold=True, size=11, space_after=4),
        make_paragraph('PBKDF2 密钥派生函数使用60万次SHA-256迭代，极大增加了暴力破解的计算成本。即使攻击者知道派生算法，单次密钥尝试也需要大量计算资源。', size=11, space_after=6),

        make_paragraph('🔐 随机化加密', bold=True, size=11, space_after=4),
        make_paragraph('每次加密使用12字节随机Nonce，确保相同密码每次加密产生不同的密文，防止通过密文模式分析推断密码信息。', size=11, space_after=6),

        make_heading('4.4.4 存储示例', level=3),
        make_paragraph('设备清单文件（device_list.txt）中的密码存储格式：', size=11, space_after=6),

        make_paragraph('加密前（明文，不安全）：', bold=True, size=10, space_after=2),
        make_paragraph('Core-SW-01|192.168.1.1|华为|核心交换机|核心层|ssh|admin|MyPassword123|', size=10, color=(134, 144, 156), space_after=6),

        make_paragraph('加密后（密文，安全）：', bold=True, size=10, space_after=2),
        make_paragraph('Core-SW-01|192.168.1.1|华为|核心交换机|核心层|ssh|admin|ENC:AbCdEf123456...|', size=10, color=(0, 180, 42), space_after=6),

        make_paragraph('密码字段以 ENC: 开头表示已加密存储，程序自动识别并解密使用，用户无需关心加解密过程。', size=11, space_after=8),

        make_heading('4.4.5 安全建议', level=3),
        make_bullet('定期更换设备密码，并在平台中更新设备清单'),
        make_bullet('设备清单文件（device_list.txt）虽然已加密，仍建议设置合适的文件访问权限'),
        make_bullet('不要将设备清单文件通过不安全的渠道（如明文邮件）传输'),
        make_bullet('更换计算机后，需要重新导入设备清单并重新输入密码（因为密钥与本机绑定）'),
        make_bullet('API Key 同样采用此加密方案存储在 ai_config.json 中'),

        make_warning('加密密钥与本机绑定。如果更换电脑或重装系统，原有的加密密码将无法解密，需要重新输入设备密码。建议在安全的环境中备份明文密码。'),

        make_paragraph(''),
    ]

    # Insert all elements before the target
    for elem in reversed(elements_to_insert):
        target_element.addprevious(elem)

    doc.save(doc_path)
    print('说明书已更新，密码安全章节已添加！')
else:
    print('未找到"常见问题"章节，请检查文档结构。')
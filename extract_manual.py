#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word文档提取手册内容并转换为文本格式
"""

import sys
from pathlib import Path
from typing import List

try:
    from docx import Document
except ImportError:
    print("错误：需要安装python-docx库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def extract_paragraphs(doc: Document) -> List[str]:
    """提取文档段落内容并格式化

    Args:
        doc: Word文档对象

    Returns:
        格式化后的文本行列表
    """
    lines: List[str] = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        text = para.text.strip()

        if not text:
            lines.append('')
            continue

        if 'Heading 1' in style:
            lines.append(f'\n{"="*60}')
            lines.append(f'  {text}')
            lines.append(f'{"="*60}')
        elif 'Heading 2' in style:
            lines.append(f'\n  {"-"*50}')
            lines.append(f'  {text}')
            lines.append(f'  {"-"*50}')
        elif 'Heading 3' in style:
            lines.append(f'\n  >> {text}')
        elif 'List Bullet' in style:
            lines.append(f'    - {text}')
        else:
            lines.append(f'  {text}')

    return lines

def extract_tables(doc: Document, lines: List[str]):
    """提取文档表格内容并格式化

    Args:
        doc: Word文档对象
        lines: 文本行列表（追加模式）
    """
    lines.append('\n\n' + '='*60)
    lines.append('  表格内容')
    lines.append('='*60)

    for i, table in enumerate(doc.tables):
        lines.append(f'\n【表格 {i+1}】')
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('  | ' + ' | '.join(cells))
            if j == 0:
                lines.append('  |' + '|'.join(['---' for _ in cells]) + '|')

def main():
    """主函数"""
    # 配置编码
    sys.stdout.reconfigure(encoding='utf-8')

    # 文件路径
    input_file = Path(r'E:\LobsterAI\Network-Config\NetOps_V1.0_使用说明书.docx')
    output_file = Path(r'E:\LobsterAI\Network-Config\manual_full.txt')

    # 验证输入文件
    if not input_file.exists():
        print(f'错误：输入文件不存在: {input_file}')
        sys.exit(1)

    # 确保输出目录存在
    output_file.parent.mkdir(parents=True, exist_ok=True)

    try:
        # 读取文档
        print(f'正在读取文档: {input_file}')
        doc = Document(str(input_file))

        # 提取内容
        lines = extract_paragraphs(doc)
        extract_tables(doc, lines)

        # 写入文件
        print(f'正在写入文件: {output_file}')
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        print('✅ 完成!')
        print(f'输出文件: {output_file}')

    except Exception as e:
        print(f'❌ 错误: {str(e)}')
        sys.exit(1)

if __name__ == '__main__':
    main()
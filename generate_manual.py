#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成NetOps使用手册（Word格式）
"""

import os
import sys
from pathlib import Path
from typing import Tuple, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：需要安装python-docx库")
    print("请运行: pip install python-docx")
    sys.exit(1)


class ManualGenerator:
    """手册生成器类"""

    def __init__(self):
        """初始化文档"""
        self.doc = Document()
        self._setup_styles()
        self._setup_sections()

    def _setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    def _setup_sections(self):
        """设置页面边距"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.8)

    def set_cell_shading(self, cell, color: Tuple[int, int, int]):
        """设置单元格背景色

        Args:
            cell: 单元格对象
            color: RGB颜色元组
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
        shading_elm.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def add_heading_styled(self, text: str, level: int = 1):
        """添加样式化的标题

        Args:
            text: 标题文本
            level: 标题级别（1-9）

        Returns:
            标题对象
        """
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        return heading

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                     size: int = 11, color: Optional[Tuple[int, int, int]] = None,
                     alignment: Optional[int] = None, space_after: int = 6):
        """添加段落

        Args:
            text: 段落文本
            bold: 是否加粗
            italic: 是否斜体
            size: 字体大小
            color: 字体颜色（RGB元组）
            alignment: 对齐方式
            space_after: 段后间距
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_after = Pt(space_after)

    def save(self, filepath: str):
        """保存文档

        Args:
            filepath: 文件路径
        """
        try:
            self.doc.save(filepath)
            print(f'✅ 文档已保存: {filepath}')
        except Exception as e:
            print(f'❌ 保存失败: {str(e)}')
            sys.exit(1)


def main():
    """主函数"""
    generator = ManualGenerator()

    # 添加标题
    generator.add_heading_styled('NetOps网络自动化运维工具使用手册', level=1)
    generator.add_paragraph('版本 1.0', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # 添加目录占位符
    generator.add_heading_styled('目录', level=2)
    generator.add_paragraph('（自动生成目录）', italic=True, color=(128, 128, 128))

    # 保存文档
    output_file = Path('NetOps使用手册.docx')
    generator.save(str(output_file))


if __name__ == '__main__':
    main()
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '165DFF')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F7FA')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_tip_box(text):
    p = doc.add_paragraph()
    run = p.add_run('💡 提示：')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(22, 93, 255)
    run2 = p.add_run(text)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(78, 89, 105)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_warning_box(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠ 注意：')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(245, 63, 63)
    run2 = p.add_run(text)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(78, 89, 105)
    p.paragraph_format.space_after = Pt(4)
    return p

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('NetOps 企业网络自动化运维平台')
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(22, 93, 255)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('V1.0 试用版 · 使用说明书')
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(78, 89, 105)

doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run('版本：V1.0 试用版\n更新日期：2026年5月\n适用平台：Windows 10/11')
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(134, 144, 156)

doc.add_page_break()

# ==================== 目录页 ====================
add_heading_styled('目录', level=1)
toc_items = [
    '1. 产品概述',
    '2. 快速入门',
    '   2.1 安装与启动',
    '   2.2 界面概览',
    '   2.3 快捷键一览',
    '3. 功能模块详解',
    '   3.1 设备配置',
    '   3.2 项目管理',
    '   3.3 运维工具箱',
    '   3.4 单点巡检',
    '   3.5 AI 智能分析',
    '   3.6 模型设置',
    '4. 典型工作流程',
    '   4.1 新项目上线流程',
    '   4.2 故障排查流程',
    '   4.3 合规审计流程',
    '   4.4 设备密码加密存储（信息安全）',
    '5. 常见问题',
    '6. 技术支持',
]
for item in toc_items:
    add_para(item, size=11, space_after=4)

doc.add_page_break()

# ==================== 1. 产品概述 ====================
add_heading_styled('1. 产品概述', level=1)

add_para('NetOps 企业网络自动化运维平台是一款面向网络工程师的多厂商网络设备配置脚本生成与自动化运维工具。平台支持锐捷、华为、华三（H3C）、思科（Cisco）四大主流厂商的网络设备，覆盖接入交换机、核心交换机、路由器、AC控制器等常见设备类型。', size=11)

add_heading_styled('1.1 核心功能', level=2)
add_bullet('多厂商设备配置脚本一键生成')
add_bullet('项目化管理，支持设备清单导入/导出')
add_bullet('批量配置备份与全网自动化巡检')
add_bullet('单点设备快速巡检与配置备份')
add_bullet('基于大模型的AI智能故障诊断与合规审计')
add_bullet('故障设备自动识别与二次核查')

add_heading_styled('1.2 支持的厂商与设备', level=2)
add_table(
    ['厂商', '接入交换机', '核心交换机', '路由器', 'AC控制器'],
    [
        ['锐捷 (Ruijie)', '✅', '✅', '✅', '✅'],
        ['华为 (Huawei)', '✅', '✅', '✅', '✅'],
        ['华三 (H3C)', '✅', '✅', '✅', '✅'],
        ['思科 (Cisco)', '✅', '✅', '✅', '✅'],
    ],
    col_widths=[3.5, 3, 3, 3, 3]
)

add_heading_styled('1.3 试用版说明', level=2)
add_para('本版本为 V1.0 试用版，面向网络工程师免费开放使用。试用版包含全部核心功能，可用于学习交流与工程实施。', size=11)
add_warning_box('试用版仅供学习交流与工程实施参考使用，不代表任何厂商官方立场，无任何官方认证。生产环境使用前请充分测试。')

doc.add_page_break()

# ==================== 2. 快速入门 ====================
add_heading_styled('2. 快速入门', level=1)

add_heading_styled('2.1 安装与启动', level=2)
add_para('本软件为绿色免安装版本，无需额外安装步骤。', size=11)
add_para('启动步骤：', bold=True, size=11)
add_bullet('双击运行 NetworkConfigGenerator.exe 即可启动程序')
add_bullet('首次启动可能需要几秒钟加载时间，请耐心等待')
add_bullet('程序启动后自动显示主界面，默认进入"设备配置"模块')

add_heading_styled('2.2 界面概览', level=2)
add_para('主界面由以下区域组成：', size=11)
add_bullet('顶部导航栏：包含 NetOps 标识、六个功能模块按钮、当前项目状态显示和"关于"按钮')
add_bullet('功能区域：占据界面主体，显示当前选中模块的操作界面')
add_bullet('底部状态栏：显示当前模块信息和快捷键提示')

add_para('六个功能模块：', bold=True, size=11)
add_table(
    ['模块名称', '图标', '快捷键', '功能说明'],
    [
        ['设备配置', '🖥', 'Ctrl+1', '选择厂商和设备类型，生成配置脚本'],
        ['新建项目', '📁', 'Ctrl+2', '创建和管理项目，维护设备清单'],
        ['运维工具箱', '🔧', 'Ctrl+3', '批量备份、全网巡检、故障核查'],
        ['单点巡检', '🔍', 'Ctrl+4', '对单台设备执行只读巡检和备份'],
        ['AI分析', '🤖', 'Ctrl+5', 'AI智能故障诊断与配置合规审计'],
        ['模型设置', '⚙', 'Ctrl+6', '配置AI大模型连接参数'],
    ],
    col_widths=[2.5, 1.5, 2, 9.5]
)

add_heading_styled('2.3 快捷键一览', level=2)
add_table(
    ['快捷键', '功能'],
    [
        ['Ctrl + 1', '切换到"设备配置"模块'],
        ['Ctrl + 2', '切换到"新建项目"模块'],
        ['Ctrl + 3', '切换到"运维工具箱"模块'],
        ['Ctrl + 4', '切换到"单点巡检"模块'],
        ['Ctrl + 5', '切换到"AI分析"模块'],
        ['Ctrl + 6', '切换到"模型设置"模块'],
    ],
    col_widths=[4, 11.5]
)

doc.add_page_break()

# ==================== 3. 功能模块详解 ====================
add_heading_styled('3. 功能模块详解', level=1)

# --- 3.1 设备配置 ---
add_heading_styled('3.1 设备配置', level=2)
add_para('设备配置模块用于生成多厂商网络设备的配置脚本。通过选择厂商和设备类型，系统自动加载对应的配置模板，用户填写参数后即可一键生成标准化的设备配置。', size=11)

add_para('操作步骤：', bold=True, size=11)
add_bullet('步骤1：在顶部选择栏点击目标厂商（锐捷/华为/华三/思科），选中后按钮变为蓝色')
add_bullet('步骤2：点击目标设备类型（接入交换机/核心交换机/路由器/AC），选中后按钮变为蓝色')
add_bullet('步骤3：系统自动加载对应的配置页面，填写设备参数（如主机名、管理IP、VLAN等）')
add_bullet('步骤4：点击"生成配置"按钮，系统自动生成标准化配置脚本')
add_bullet('步骤5：可复制生成的配置脚本，或保存到项目目录中')

add_tip_box('必须先选择厂商，再选择设备类型，两者都选中后才会显示配置页面。')

add_para('支持的配置类型：', bold=True, size=11)
add_table(
    ['设备类型', '典型配置内容'],
    [
        ['接入交换机', 'VLAN配置、端口安全、STP、AAA认证、SNMP、ACL等'],
        ['核心交换机', 'VLAN接口、路由协议(OSPF/BGP)、VRRP/HSRP、链路聚合等'],
        ['路由器', '接口配置、路由协议、NAT、VPN、QoS、ACL等'],
        ['AC控制器', '无线网络配置、SSID、安全策略、AP管理、射频优化等'],
    ],
    col_widths=[3.5, 12]
)

# --- 3.2 项目管理 ---
add_heading_styled('3.2 项目管理', level=2)
add_para('项目管理模块是整个平台的核心组织单元。每个项目对应一个独立的网络工程项目，包含设备清单、配置脚本、巡检报告、配置备份等完整数据。', size=11)

add_heading_styled('3.2.1 创建项目', level=3)
add_para('操作步骤：', bold=True, size=11)
add_bullet('在"新建项目"区域输入项目名称（如：栖霞区电子政务外网项目）')
add_bullet('点击"创建项目"按钮')
add_bullet('系统自动在 C:\\Network-Config\\projects\\ 目录下创建项目文件夹')
add_bullet('项目文件夹包含以下子目录：config（配置）、output（输出）、report（报告）、config_backup（配置备份）')

add_heading_styled('3.2.2 切换与管理项目', level=3)
add_para('在"项目切换"区域可以进行以下操作：', size=11)
add_bullet('切换项目：在下拉框中选择目标项目，点击"切换项目"按钮')
add_bullet('刷新列表：点击"刷新列表"按钮更新项目下拉框')
add_bullet('编辑项目：选中项目后点击"编辑项目"可修改项目名称')
add_bullet('删除项目：选中项目后点击"删除项目"可删除整个项目（含所有数据，操作不可逆）')

add_warning_box('删除项目将永久删除项目文件夹及其所有数据，请谨慎操作！')

add_heading_styled('3.2.3 设备清单管理', level=3)
add_para('设备清单是运维自动化的数据基础，包含项目中所有网络设备的连接信息。', size=11)

add_para('设备清单包含以下字段：', bold=True, size=11)
add_table(
    ['字段', '说明', '示例'],
    [
        ['设备名称', '设备的标识名称', 'Core-SW-01'],
        ['管理IP', '设备的管理地址', '192.168.1.1'],
        ['厂商', '设备品牌', '华为 / 锐捷 / H3C / 思科'],
        ['设备类型', '设备角色', '核心交换机 / 路由器 / AC'],
        ['分组', '自定义分组标签', '核心层 / 汇聚层 / 接入层'],
        ['协议', '连接协议', 'ssh / telnet'],
        ['用户名', 'SSH/Telnet登录用户名', 'admin'],
        ['密码', 'SSH/Telnet登录密码', '********'],
        ['特权密码', 'Cisco设备enable密码（可选）', '********'],
    ],
    col_widths=[2.5, 5, 8]
)

add_para('设备清单操作按钮：', bold=True, size=11)
add_table(
    ['按钮', '功能'],
    [
        ['新增设备', '打开设备信息表单，逐台添加设备'],
        ['编辑设备', '修改选中设备的各项参数'],
        ['删除选中', '从清单中移除选中的设备行'],
        ['保存清单', '将当前设备清单保存到项目config目录'],
        ['导入', '从CSV/Excel文件批量导入设备清单'],
        ['导出', '将当前设备清单导出为CSV文件'],
        ['下载模板', '下载标准设备清单模板文件'],
        ['模板库', '打开预设的设备模板库，快速填充'],
        ['批量发现', '通过IP扫描自动发现网络中的设备'],
        ['连接测试', '测试选中设备的SSH/Telnet连通性'],
        ['变更历史', '查看设备清单的修改历史记录'],
    ],
    col_widths=[3, 12.5]
)

add_tip_box('建议先使用"下载模板"获取标准格式文件，按模板填写设备信息后使用"导入"功能批量导入，效率最高。')

# --- 3.3 运维工具箱 ---
add_heading_styled('3.3 运维工具箱', level=2)
add_para('运维工具箱提供三个独立的自动化运维任务，每个任务可独立选择项目，互不干扰。所有操作仅执行只读指令，不会修改设备任何配置。', size=11)

add_heading_styled('3.3.1 批量配置备份', level=3)
add_para('功能说明：读取项目设备清单，批量SSH登录所有设备，抓取运行配置（running-config），按"设备IP+设备类型"命名存入项目的 config_backup 目录。', size=11)
add_para('操作步骤：', bold=True, size=11)
add_bullet('在"批量配置备份"面板中选择目标项目')
add_bullet('点击"执行批量备份"按钮')
add_bullet('观察进度条和日志输出，等待任务完成')
add_bullet('备份文件保存在 项目目录\\config_backup\\ 下')

add_heading_styled('3.3.2 全网自动化巡检', level=3)
add_para('功能说明：对项目中的所有设备执行CPU、内存、接口状态、日志等只读巡检指令，生成全网巡检总报告，自动识别故障设备并单独归档。', size=11)
add_para('操作步骤：', bold=True, size=11)
add_bullet('在"全网自动化巡检"面板中选择目标项目')
add_bullet('点击"执行全网巡检"按钮')
add_bullet('系统自动逐台登录设备执行巡检命令')
add_bullet('巡检完成后，报告保存在 项目目录\\output\\ 下')
add_bullet('异常设备信息单独归档到 output\\single_exception\\ 目录')

add_heading_styled('3.3.3 故障设备二次核查', level=3)
add_para('功能说明：自动读取上一轮巡检中识别的故障设备IP，执行合规只读指令进行二次核查。高危配置指令自动剥离，生成人工操作清单，确保安全。', size=11)
add_para('操作步骤：', bold=True, size=11)
add_bullet('在"故障设备二次核查"面板中选择目标项目')
add_bullet('点击"执行故障核查"按钮')
add_bullet('系统仅对故障设备执行核查，节省时间')
add_bullet('核查结果保存在 项目目录\\output\\trouble_check_result\\ 下')

add_warning_box('所有运维任务仅执行只读命令（show、display等），不会修改设备配置。但请确保使用只读权限的账号。')

# --- 3.4 单点巡检 ---
add_heading_styled('3.4 单点巡检', level=2)
add_para('单点巡检模块不依赖项目设备清单，可直接输入单台设备的IP和凭证，执行只读巡检。支持一键备份配置并联动AI分析。适用于临时排查、单台设备验收等场景。', size=11)

add_para('操作步骤：', bold=True, size=11)
add_bullet('步骤1：填写设备连接信息（IP地址、协议、厂商、设备类型）')
add_bullet('步骤2：填写登录凭证（用户名、密码，Cisco设备需填写特权密码）')
add_bullet('步骤3：选择执行操作：')
add_bullet('单设备巡检：执行只读巡检命令，输出设备健康状态报告', level=1)
add_bullet('单设备备份：SSH登录抓取运行配置并保存', level=1)
add_bullet('AI分析结果：将巡检结果发送给AI模型进行智能分析', level=1)
add_bullet('步骤4：观察进度条和状态提示，等待任务完成')

add_tip_box('单点巡检适合快速排查单台设备问题，无需创建项目即可使用。巡检结果可联动AI分析模块进行深度诊断。')

# --- 3.5 AI 智能分析 ---
add_heading_styled('3.5 AI 智能分析', level=2)
add_para('AI智能分析模块基于大语言模型，对设备巡检日志进行智能故障诊断，或对设备配置进行合规性审计。使用前请先在"模型设置"中配置AI模型连接。', size=11)

add_heading_styled('3.5.1 界面布局', level=3)
add_para('AI分析页面包含以下交互区域：', size=11)
add_bullet('当前项目：下拉框选择项目，支持右键菜单刷新/清除，右侧✕按钮可清除选择')
add_bullet('分析模式：下拉框选择"故障诊断"或"配置合规审计"')
add_bullet('目标文件：显示当前选中的文件路径，支持浏览选择、清除、最近文件记录')
add_bullet('批量模式：支持同时选择多个文件进行分析')
add_bullet('输入数据区：支持粘贴文本、拖拽文件（.txt/.log/.md）、从剪贴板粘贴、加载文件')
add_bullet('文件预览：自动显示加载文件的前5行内容，方便确认文件正确性')
add_bullet('分析结果区：显示AI返回的结构化分析报告，支持保存和导出')

add_heading_styled('3.5.2 两种分析模式', level=3)
add_table(
    ['分析模式', '输入内容', '输出内容'],
    [
        ['故障诊断', '设备巡检日志（show命令输出）', '异常指标识别、严重程度分级、原因分析、修复建议、健康度评分'],
        ['配置合规审计', '设备配置文件', '安全配置检查、协议合规性、冗余检查、风险分级、整改建议'],
    ],
    col_widths=[3, 5.5, 7]
)

add_heading_styled('3.5.3 操作步骤', level=3)
add_bullet('步骤1：在"模型设置"中配置好AI模型（首次使用必须）')
add_bullet('步骤2：选择分析模式（故障诊断 / 配置合规审计）')
add_bullet('步骤3：加载分析数据（三种方式任选）：')
add_bullet('方式A：点击"浏览"选择日志/配置文件', level=1)
add_bullet('方式B：直接拖拽文件到输入区域', level=1)
add_bullet('方式C：粘贴文本内容到输入区域', level=1)
add_bullet('步骤4：确认文件预览内容正确')
add_bullet('步骤5：点击"开始分析"按钮')
add_bullet('步骤6：等待AI返回分析结果（可能需要数十秒）')
add_bullet('步骤7：查看分析报告，可点击"保存报告"或"导出TXT"保存结果')

add_tip_box('最近使用的文件会自动记录在"最近文件"下拉框中（最多5个），方便快速切换。支持批量选择多个文件逐一分析。')

# --- 3.6 模型设置 ---
add_heading_styled('3.6 模型设置', level=2)
add_para('模型设置模块用于配置AI大模型的连接参数，是使用AI分析功能的前提条件。', size=11)

add_heading_styled('3.6.1 支持的AI模型厂商', level=3)
add_table(
    ['厂商', '默认BaseURL', '默认模型', '需要API Key'],
    [
        ['OpenAI', 'https://api.openai.com/v1', 'gpt-4o', '是'],
        ['DeepSeek', 'https://api.deepseek.com/v1', 'deepseek-chat', '是'],
        ['通义千问', 'https://dashscope.aliyuncs.com/compatible-mode/v1', 'qwen-plus', '是'],
        ['智谱GLM', 'https://open.bigmodel.cn/api/paas/v4', 'glm-4', '是'],
        ['Moonshot', 'https://api.moonshot.cn/v1', 'moonshot-v1-8k', '是'],
        ['自定义', '用户自行填写', '用户自行填写', '是'],
    ],
    col_widths=[2.5, 6, 3.5, 3.5]
)

add_heading_styled('3.6.2 配置步骤', level=3)
add_bullet('步骤1：在"模型厂商"下拉框中选择AI服务商（或选择"自定义"）')
add_bullet('步骤2：系统自动填充默认的 BaseURL 和模型名称（自定义模式需手动填写）')
add_bullet('步骤3：输入您的 API Key（从AI服务商官网获取）')
add_bullet('步骤4：点击"测试连接"按钮验证配置是否正确')
add_bullet('步骤5：测试成功后点击"保存配置"')

add_tip_box('选择预设厂商后，BaseURL和模型名称会自动填充，只需填写API Key即可。API Key请从对应厂商的官方平台获取。')

add_warning_box('API Key 是敏感信息，请妥善保管。本软件将API Key加密存储在本地配置文件中，不会上传到任何服务器。')

doc.add_page_break()

# ==================== 4. 典型工作流程 ====================
add_heading_styled('4. 典型工作流程', level=1)

add_heading_styled('4.1 新项目上线流程', level=2)
add_para('以下是一个典型的新网络工程项目使用本平台的完整流程：', size=11)
add_bullet('第1步：创建项目 → 在"新建项目"模块输入项目名称并创建')
add_bullet('第2步：导入设备 → 下载模板，填写设备清单，导入到项目中')
add_bullet('第3步：生成配置 → 在"设备配置"模块选择厂商和类型，生成各设备配置脚本')
add_bullet('第4步：部署设备 → 将生成的配置脚本部署到实际设备上')
add_bullet('第5步：配置备份 → 在"运维工具箱"执行批量配置备份，存档运行配置')
add_bullet('第6步：全网巡检 → 执行全网自动化巡检，验证设备运行状态')

add_heading_styled('4.2 故障排查流程', level=2)
add_bullet('第1步：单点巡检 → 在"单点巡检"模块输入故障设备IP，执行快速巡检')
add_bullet('第2步：AI诊断 → 将巡检结果粘贴到"AI分析"模块，选择"故障诊断"模式')
add_bullet('第3步：查看报告 → AI输出结构化的故障诊断报告，包含原因分析和修复建议')
add_bullet('第4步：全网核查 → 如怀疑多台设备受影响，在"运维工具箱"执行全网巡检')
add_bullet('第5步：二次确认 → 对故障设备执行"故障设备二次核查"')

add_heading_styled('4.3 合规审计流程', level=2)
add_bullet('第1步：获取配置 → 在"运维工具箱"执行批量配置备份，获取所有设备运行配置')
add_bullet('第2步：AI审计 → 在"AI分析"模块选择"配置合规审计"模式，加载配置文件')
add_bullet('第3步：整改 → 根据AI审计报告中的建议，逐项整改配置风险')
add_bullet('第4步：复检 → 整改完成后再次执行审计，确认所有风险项已处理')

# --- 4.4 设备密码加密存储 ---
add_heading_styled('4.4 设备密码加密存储（信息安全）', level=2)
add_para('本平台高度重视信息安全，所有设备登录密码均采用业界标准的加密方案进行存储，确保敏感数据不会以明文形式泄露。', size=11)

add_heading_styled('4.4.1 加密算法概述', level=3)
add_para('设备密码加密采用 AES-256-GCM 认证加密算法，这是目前业界公认最安全的对称加密方案之一，被广泛应用于金融、政府和军事领域。', size=11)

add_table(
    ['加密组件', '技术规格', '安全说明'],
    [
        ['加密算法', 'AES-256-GCM', '256位密钥，GCM模式提供认证加密，同时保证机密性和完整性'],
        ['密钥派生', 'PBKDF2-HMAC-SHA256', '基于密码的密钥派生函数，使用SHA-256哈希，迭代60万次，有效抵御暴力破解'],
        ['密钥来源', '本机唯一标识', '基于Windows MachineGuid + 主机名生成，每台计算机的密钥唯一'],
        ['随机数(Nonce)', '12字节随机生成', '每次加密使用不同的随机Nonce，相同密码每次加密结果不同'],
        ['存储格式', 'ENC: + Base64编码', '密文以 ENC: 前缀标识，Base64编码存储，便于识别和传输'],
        ['加密库', 'Python Cryptography', '使用Python官方推荐的cryptography库，经过广泛安全审计'],
    ],
    col_widths=[3, 4.5, 8]
)

add_heading_styled('4.4.2 加密流程', level=3)
add_para('密码加密与解密流程如下：', size=11)

add_para('加密过程（保存密码时）：', bold=True, size=11)
add_bullet('① 采集本机唯一标识（Windows MachineGuid + 主机名）')
add_bullet('② 通过 PBKDF2-HMAC-SHA256 派生256位加密密钥（迭代60万次）')
add_bullet('③ 生成12字节随机 Nonce')
add_bullet('④ 使用 AES-256-GCM 加密明文密码')
add_bullet('⑤ 将 Nonce + 密文 进行 Base64 编码')
add_bullet('⑥ 添加 ENC: 前缀标识，写入设备清单文件')

add_para('解密过程（使用密码时）：', bold=True, size=11)
add_bullet('① 检测密码字段是否以 ENC: 开头')
add_bullet('② 从本机标识重新派生相同的加密密钥')
add_bullet('③ Base64解码，分离 Nonce 和密文')
add_bullet('④ 使用 AES-256-GCM 解密并验证完整性')
add_bullet('⑤ 返回明文密码供SSH/Telnet连接使用')

add_heading_styled('4.4.3 安全特性', level=3)

add_para('🔐 本机绑定', bold=True, size=11)
add_para('加密密钥由本机唯一标识派生，即使他人获取了加密后的设备清单文件，也无法在其他计算机上解密。这确保了：', size=11)
add_bullet('设备清单文件被拷贝到其他电脑后，密码无法解密')
add_bullet('每台计算机的加密密钥独立，互不通用')
add_bullet('有效防止文件泄露导致的密码外泄')

add_para('🔐 认证加密', bold=True, size=11)
add_para('AES-GCM 模式不仅加密数据，还提供完整性校验。任何对密文的篡改都会导致解密失败并被检测到，防止密文被恶意修改。', size=11)

add_para('🔐 抗暴力破解', bold=True, size=11)
add_para('PBKDF2 密钥派生函数使用60万次SHA-256迭代，极大增加了暴力破解的计算成本。即使攻击者知道派生算法，单次密钥尝试也需要大量计算资源。', size=11)

add_para('🔐 随机化加密', bold=True, size=11)
add_para('每次加密使用12字节随机Nonce，确保相同密码每次加密产生不同的密文，防止通过密文模式分析推断密码信息。', size=11)

add_heading_styled('4.4.4 存储示例', level=3)
add_para('设备清单文件（device_list.txt）中的密码存储格式：', size=11)

add_para('加密前（明文，不安全）：', bold=True, size=10, color=(134, 144, 156))
add_para('Core-SW-01|192.168.1.1|华为|核心交换机|核心层|ssh|admin|MyPassword123|', size=10, color=(134, 144, 156))

add_para('加密后（密文，安全）：', bold=True, size=10, color=(0, 180, 42))
add_para('Core-SW-01|192.168.1.1|华为|核心交换机|核心层|ssh|admin|ENC:AbCdEf123456...|', size=10, color=(0, 180, 42))

add_para('密码字段以 ENC: 开头表示已加密存储，程序自动识别并解密使用，用户无需关心加解密过程。', size=11)

add_heading_styled('4.4.5 安全建议', level=3)
add_bullet('定期更换设备密码，并在平台中更新设备清单')
add_bullet('设备清单文件（device_list.txt）虽然已加密，仍建议设置合适的文件访问权限')
add_bullet('不要将设备清单文件通过不安全的渠道（如明文邮件）传输')
add_bullet('更换计算机后，需要重新导入设备清单并重新输入密码（因为密钥与本机绑定）')
add_bullet('API Key 同样采用此加密方案存储在 ai_config.json 中')

add_warning_box('加密密钥与本机绑定。如果更换电脑或重装系统，原有的加密密码将无法解密，需要重新输入设备密码。建议在安全的环境中备份明文密码。')

doc.add_page_break()

# ==================== 5. 常见问题 ====================
add_heading_styled('5. 常见问题', level=1)

faqs = [
    ('Q1：程序启动后界面显示异常怎么办？',
     '请确认您的操作系统为 Windows 10 或 Windows 11，屏幕分辨率不低于 1366×768。如仍有问题，请尝试以管理员身份运行。'),
    ('Q2：SSH连接设备失败怎么办？',
     '请检查：① 设备IP地址是否正确；② 设备是否开启了SSH服务；③ 用户名和密码是否正确；④ 网络是否可达（可使用"连接测试"功能验证）。'),
    ('Q3：AI分析功能无法使用？',
     '请确认已在"模型设置"中正确配置了AI模型参数，并通过了"测试连接"。需要有效的API Key和网络连接。'),
    ('Q4：如何批量导入设备清单？',
     '点击"下载模板"获取标准CSV模板，按模板格式填写设备信息后，点击"导入"按钮选择文件即可。'),
    ('Q5：项目删除后能恢复吗？',
     '项目删除是永久性的，会删除项目文件夹及其所有数据。建议删除前先确认项目数据已备份。'),
    ('Q6：运维任务会修改设备配置吗？',
     '不会。所有运维任务仅执行只读命令（show、display等），不会对设备配置做任何修改。请使用只读权限的账号以确保安全。'),
    ('Q7：支持哪些设备型号？',
     '理论上支持所有支持SSH/Telnet的标准网络设备。已测试的厂商包括锐捷、华为、H3C、思科的主流交换机和路由器。'),
    ('Q8：如何获取API Key？',
     '请访问对应AI厂商的官方网站注册并获取API Key。例如：OpenAI用户访问 platform.openai.com，DeepSeek用户访问 platform.deepseek.com。'),
]

for q, a in faqs:
    add_para(q, bold=True, size=11, space_after=2)
    add_para(a, size=11, space_after=8)

doc.add_page_break()

# ==================== 6. 技术支持 ====================
add_heading_styled('6. 技术支持', level=1)

add_para('版本信息', bold=True, size=12)
add_table(
    ['项目', '信息'],
    [
        ['软件名称', 'NetOps 企业网络自动化运维平台'],
        ['版本号', 'V1.0 试用版'],
        ['适用系统', 'Windows 10 / Windows 11'],
        ['开发语言', 'Python 3.11 + PyQt5'],
        ['开源协议', 'MIT License'],
        ['项目地址', 'https://github.com/hanchunjun/network-config-generator'],
    ],
    col_widths=[3, 12.5]
)

add_para('')
add_para('免责声明', bold=True, size=12)
add_para('本软件源码基于 MIT License 开源发布，可自由获取与修改。软件全功能使用需授权激活。使用者应自行评估软件适用性，并对使用结果负责。本软件不代表任何厂商官方立场，无任何官方认证。在生产环境中使用前，请务必在测试环境中充分验证。', size=11)

add_para('')
add_para('联系方式', bold=True, size=12)
add_para('如有问题或建议，欢迎通过以下方式联系：', size=11)
add_bullet('GitHub Issues：https://github.com/hanchunjun/network-config-generator/issues')
add_bullet('项目主页：https://github.com/hanchunjun/network-config-generator')

add_para('')
add_para('Copyright © 2026 laohan. All rights reserved.', size=10, color=(134, 144, 156), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ==================== 保存 ====================
output_path = os.path.join(os.path.dirname(__file__), 'NetOps_V1.0_使用说明书.docx')
doc.save(output_path)
print(f'使用说明书已生成：{output_path}')